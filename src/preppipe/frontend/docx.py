# SPDX-FileCopyrightText: 2023 PrepPipe's Contributors
# SPDX-License-Identifier: Apache-2.0

import docx
import docx.document
import docx.parts
import docx.parts.image
import docx.text.run
import docx.text.paragraph
import docx.enum.text
import docx.styles.style
import docx.table
import docx.oxml.text.paragraph
import docx.oxml.table
import lxml
import lxml.etree
import zipfile
import os
import sys
import typing
import dataclasses

from ..irbase import *
from ..inputmodel import *
from ..pipeline import *

class _DOCXParseContext:
  @dataclasses.dataclass
  class CharacterStyle:
    bgcolor : Color | None = None
    fgcolor : Color | None = None
    bold : bool = False
    italic : bool = False
    strikethrough : bool = False

  @dataclasses.dataclass
  class ParagraphStyle:
    bgcolor : Color | None = None
    align_mid : bool = False
    list_level : int | None = None

  context : Context
  filepath : str
  dochandle : docx.document.Document
  ziphandle : zipfile.ZipFile
  difile : DIFile
  documentname : str
  nsprefix : typing.Literal[r"{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"]

  rels : dict[str, str]
  textstyles : dict[str, TextStyleLiteral | None]
  character_styles : dict[str, CharacterStyle]
  paragraph_styles : dict[str, ParagraphStyle]
  result : IMDocumentOp

  def __init__(self, ctx : Context, filepath : str) -> None:
    self.context = ctx
    self.filepath = filepath
    self.dochandle = docx.Document(filepath)
    self.ziphandle = zipfile.ZipFile(filepath, mode = "r")
    self.difile = ctx.get_DIFile(filepath)
    self.documentname = os.path.splitext(os.path.basename(filepath))[0]
    self.nsprefix=r"{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

    self.rels = {}
    self.textstyles = {}
    self.character_styles = {}
    self.paragraph_styles = {}
    self.result = IMDocumentOp.create(self.documentname, self.difile)

  @staticmethod
  def parse_docx(ctx : Context, filepath : str):
    _docx = _DOCXParseContext(ctx, filepath)
    _docx.parse()
    return _docx.result

  def cleanup(self):
    self.ziphandle.close()

  def parse_rels(self):
    # https://stackoverflow.com/questions/27691678/finding-image-present-docx-file-using-python/61331396#61331396?newreg=ed63b33ba36847ccbdc19f5d20494cab
    for r in self.dochandle.part.rels.values():
      if isinstance(r._target, docx.parts.image.ImagePart):
        self.rels[r.rId] = os.path.basename(r._target.partname)

  def get_paragraph_bgcolor_impl(self, ps : docx.styles.style._ParagraphStyle | docx.text.paragraph.Paragraph) -> Color | None:
    element = ps.paragraph_format.element
    # 查找 <w:pPr> 下是否有 <w:shd> 项，有的话其中的 fill 属性就是背景颜色
    # 不能用 := 因为返回的元素即使有也会成为 False
    pPr = element.find(self.nsprefix + "pPr")
    if pPr is not None:
      shd = pPr.find(self.nsprefix + "shd")
      if shd is not None:
        colorstr = shd.get(self.nsprefix + 'fill')
        if isinstance(colorstr, str):
          return Color.get('#' + colorstr)
    # 没有的话就没有了
    return None

  def get_paragraph_align_mid_impl(self, ps : docx.styles.style._ParagraphStyle | docx.text.paragraph.Paragraph) -> bool | None:
    if isinstance(ps, docx.styles.style._ParagraphStyle):
      alignment = ps.paragraph_format.alignment
    else:
      alignment = ps.alignment
    if alignment:
      return alignment == docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    return None

  def get_paragraph_listlevel_impl(self, ps: docx.styles.style._ParagraphStyle | docx.text.paragraph.Paragraph) -> int | None:
    if isinstance(ps, docx.styles.style._ParagraphStyle):
      return None
    # 查找 <w:pPr> 下是否有 <w:numPr>，有的话看里面是否有 <w:ilvl>
    # 不能用 := 因为返回的元素即使有也会成为 False
    pPr = ps.paragraph_format.element.find(self.nsprefix + "pPr")
    if pPr is not None:
      numPr = pPr.find(self.nsprefix + "numPr")
      if numPr is not None:
        ilvl = numPr.find(self.nsprefix + "ilvl")
        if ilvl is not None:
          if lvlstr := ilvl.get(self.nsprefix+"val"):
            return int(lvlstr)
    return None

  def get_character_bold_impl(self, cs : docx.styles.style._CharacterStyle | docx.text.run.Run) -> bool | None:
    return cs.font.bold

  def get_character_italic_impl(self, cs : docx.styles.style._CharacterStyle | docx.text.run.Run) -> bool | None:
    return cs.font.italic

  def get_character_strikethrough_impl(self, cs : docx.styles.style._CharacterStyle | docx.text.run.Run) -> bool | None:
    return cs.font.strike

  def get_character_fg_color_impl(self, cs : docx.styles.style._CharacterStyle | docx.text.run.Run) -> Color | None:
    if color := cs.font.color:
      if rgbcolor := color.rgb:
        return Color(r=rgbcolor[0], g=rgbcolor[1], b=rgbcolor[2])
    return None

  def get_character_bg_color_impl(self, cs : docx.styles.style._CharacterStyle | docx.text.run.Run) -> Color | None:
    if color := cs.font.highlight_color:
      match color:
        case docx.enum.text.WD_COLOR_INDEX.AUTO:
          return None
        case docx.enum.text.WD_COLOR_INDEX.BLACK:
          return Color(0,0,0)
        case docx.enum.text.WD_COLOR_INDEX.BLUE:
          return Color(0,0,255)
        case docx.enum.text.WD_COLOR_INDEX.BRIGHT_GREEN:
          return Color(170, 255, 0)
        case docx.enum.text.WD_COLOR_INDEX.DARK_BLUE:
          return Color(0,0,139)
        case docx.enum.text.WD_COLOR_INDEX.DARK_RED:
          return Color(139,0,0)
        case docx.enum.text.WD_COLOR_INDEX.DARK_YELLOW:
          return Color(139, 128, 0)
        case docx.enum.text.WD_COLOR_INDEX.GRAY_25:
          return Color(64,64,64)
        case docx.enum.text.WD_COLOR_INDEX.GRAY_50:
          return Color(127,127,127)
        case docx.enum.text.WD_COLOR_INDEX.GREEN:
          return Color(0, 255, 0)
        case docx.enum.text.WD_COLOR_INDEX.PINK:
          return Color(255,192,203)
        case docx.enum.text.WD_COLOR_INDEX.RED:
          return Color(255,0,0)
        case docx.enum.text.WD_COLOR_INDEX.TEAL:
          return Color(0,128,128)
        case docx.enum.text.WD_COLOR_INDEX.TURQUOISE:
          return Color(64,224,208)
        case docx.enum.text.WD_COLOR_INDEX.VIOLET:
          return Color(128,0,128)
        case docx.enum.text.WD_COLOR_INDEX.WHITE:
          return Color(255,255,255)
        case docx.enum.text.WD_COLOR_INDEX.YELLOW:
          return Color(255,255,0)
        case _:
          return None

  def parse_styles(self):
    for s in self.dochandle.styles:
      if s.name is None:
        continue
      if isinstance(s, docx.styles.style._CharacterStyle):
        bgcolor = self.get_character_bg_color_impl(s)
        fgcolor = self.get_character_fg_color_impl(s)
        bold = self.get_character_bold_impl(s)
        italic = self.get_character_italic_impl(s)
        strikethrough = self.get_character_strikethrough_impl(s)
        if bold is None:
          bold = False
        if italic is None:
          italic = False
        if strikethrough is None:
          strikethrough = False
        self.character_styles[s.name] = self.CharacterStyle(bgcolor=bgcolor, fgcolor=fgcolor, bold=bold, italic=italic, strikethrough=strikethrough)
      elif isinstance(s, docx.styles.style._ParagraphStyle):
        bgcolor = self.get_paragraph_bgcolor_impl(s)
        align_mid = self.get_paragraph_align_mid_impl(s)
        list_level = self.get_paragraph_listlevel_impl(s)
        if align_mid is None:
          align_mid = False
        self.paragraph_styles[s.name] = self.ParagraphStyle(bgcolor=bgcolor, align_mid=align_mid, list_level=list_level)

  def parse(self):
    self.parse_rels()
    self.parse_styles()
    self.parse_mainloop()
    self.cleanup()

  def get_path_from_rid(self, rid : str) -> str | None:
    if rid in self.rels:
      return "word/media/" + self.rels[rid]

  def get_full_path_from_epath(self, path : str) -> str:
    return os.path.normpath(os.path.join(self.filepath, path))

  def query_style(self, entity, cb : typing.Callable, stylecb : typing.Callable | None = None):
    if value := cb(entity):
      return value
    style = stylecb(entity) if stylecb is not None else entity.style
    value = cb(style)
    while value is None and style.base_style is not None:
      style = style.base_style
      value = cb(style)
    return value

  def get_paragraph_style(self, p : docx.text.paragraph.Paragraph) -> ParagraphStyle:
    bgcolor = None
    align_mid = False
    list_level = None
    if src := self.paragraph_styles.get(p.style.name, None):
      bgcolor = src.bgcolor
      align_mid = src.align_mid
      list_level = src.list_level
    if bgcolor_override := self.get_paragraph_bgcolor_impl(p):
      bgcolor = bgcolor_override

    align_mid_override = self.get_paragraph_align_mid_impl(p)
    if align_mid_override is not None:
      align_mid = align_mid_override

    list_level_override = self.get_paragraph_listlevel_impl(p)
    if list_level_override is not None:
      list_level = list_level_override
    return self.ParagraphStyle(bgcolor=bgcolor, align_mid=align_mid, list_level=list_level)

  def get_character_style(self, r : docx.text.run.Run) -> TextStyleLiteral | bool | None:
    src = None
    if r.style is not None:
      src = self.character_styles.get(r.style.name, None)
    styledict = {}

    if strikethrough := self.get_character_strikethrough_impl(r):
      return False
    elif src is not None and src.strikethrough:
      return False

    if bgcolor := self.get_character_bg_color_impl(r):
      styledict[TextAttribute.BackgroundColor] = bgcolor
    elif src is not None and src.bgcolor:
      styledict[TextAttribute.BackgroundColor] = src.bgcolor

    if fgcolor := self.get_character_fg_color_impl(r):
      styledict[TextAttribute.TextColor] = fgcolor
    elif src is not None and src.fgcolor:
      styledict[TextAttribute.TextColor] = src.fgcolor

    bold = self.get_character_bold_impl(r)
    if bold is None and src is not None:
      bold = src.bold
    if bold:
      styledict[TextAttribute.Bold] = True

    italic = self.get_character_italic_impl(r)
    if italic is None and src is not None:
      italic = src.italic
    if italic:
      styledict[TextAttribute.Italic] = True

    if len(styledict) == 0:
      return None
    return TextStyleLiteral.get(styledict, self.context)

  def get_paragraph_alignment(self, p : docx.text.paragraph.Paragraph) -> docx.enum.text.WD_PARAGRAPH_ALIGNMENT | None:
    return self.query_style(p, lambda p : p.alignment)

  def iter_block_items(self, parent):
    """
    Reference: https://github.com/python-openxml/python-docx/issues/40#issuecomment-90710401

    Yield each paragraph and table child within *parent*, in document order.
    Each returned value is an instance of either Table or Paragraph. *parent*
    would most commonly be a reference to a main Document object, but
    also works for a _Cell object, which itself can contain paragraphs and tables.
    """
    if isinstance(parent, docx.document.Document):
      parent_elm = parent._body._element
    elif isinstance(parent, docx.table._Cell):
      parent_elm = parent._tc
    else:
      raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
      if isinstance(child, docx.oxml.text.paragraph.CT_P):
        yield docx.text.paragraph.Paragraph(child, parent)
      elif isinstance(child, docx.oxml.table.CT_Tbl):
        yield docx.table.Table(child, parent)

  def parse_mainloop(self):
    # 目前暂不支持页码
    page = 0
    row = 0

    dest = self.result.body

    # 生成时直接合并相邻的特殊块和列表层级
    last_special_block : IMSpecialBlockOp | None = None
    current_ongoing_list : list[IMListOp] = []
    # 如果当前段落开始时，我们正在一个列表中，那么列表的每个层级（包括第零层）都在 current_ongoing_list 中
    # 如果当前段落开始时，前面是一个在相同列表层级、相同成因的特殊块，那么 last_special_block 指向该特殊块，我们可以直接合并内容
    def handle_paragraph(p : docx.text.paragraph.Paragraph):
      nonlocal page
      nonlocal row
      nonlocal last_special_block
      nonlocal current_ongoing_list

      row += 1
      col = 1
      # 首先检查段落样式，是否是特殊块或者是列表
      pstyle = self.get_paragraph_style(p)
      special_block_reason : str | None = None
      list_level : int | None = pstyle.list_level
      if pstyle.align_mid:
        special_block_reason = IMSpecialBlockOp.ATTR_REASON_CENTERED
      elif pstyle.bgcolor is not None and pstyle.bgcolor.to_tuple() != (255, 255, 255):
        special_block_reason = IMSpecialBlockOp.ATTR_REASON_BG_HIGHLIGHT
      # 更新现在的状态
      # current_dest_block 只在我们不可能合并特殊块时非 None
      current_dest_block : Block | None = None
      is_continuous = False
      # 如果列表层级不一致，我们重新开始
      if list_level is None and len(current_ongoing_list) > 0:
        # 一个列表结束了，即使前面有特殊块，我们也不可能合并的
        # 我们要在原处加块
        current_dest_block = dest.create_block()
        current_ongoing_list.clear()
        last_special_block = None
      elif list_level is not None and list_level == len(current_ongoing_list) - 1:
        # 我们正在沿用前一个层级的内容
        is_continuous = True
      elif list_level is not None:
        # 我们现在在一个列表中，但是层级与前面的不一致
        # 如果层级不够就补足，层级多了就出栈
        # 不过不管怎样，特殊块都不能合并了
        last_special_block = None
        while list_level < len(current_ongoing_list) -1:
          current_ongoing_list.pop()
        while list_level >= len(current_ongoing_list):
          # 我们需要补足列表层级
          curlistlevel = IMListOp.create('', self.context.get_DILocation(self.difile, page, row, col))
          if len(current_ongoing_list) == 0:
            # 这是列表的第一级
            parentblock = dest.create_block()
          else:
            # 这是某个列表下的一级
            lastitem = current_ongoing_list[-1].get_last_region()
            if lastitem is None:
              lastitem = current_ongoing_list[-1].add_list_item()
            parentblock = lastitem.create_block()
          parentblock.push_back(curlistlevel)
          current_ongoing_list.append(curlistlevel)
        current_dest_block = current_ongoing_list[-1].add_list_item().create_block()
      else:
        # 前面也没有列表，我们也没有开始新的列表
        pass

      # 列表层级判断完了
      # 看看我们是否可能合并特殊块
      if last_special_block is None or special_block_reason is None:
        if current_dest_block is None:
          # 我们不会需要合并特殊块，需要在此创建目标块
          if len(current_ongoing_list) > 0:
            current_dest_block = current_ongoing_list[-1].add_list_item().create_block()
          else:
            current_dest_block = dest.create_block()

      # 然后我们把内容弄出来
      # 即使是特殊块，我们也要按照正常的读取方式先读取一遍，这样可以不漏掉错误和内嵌的图片
      # 如果我们要续特殊块的话，也需要在没有其他
      pending_contents = _DocxTextCoalescer(self.context)
      for r in p.runs:
        res = self.parse_run(r, page, row, col)
        if res is not None:
          pending_contents.add_value(res, col)
        # 不管怎样，更新列号
        # 如果字体是被划去的，那么 res 也是 None，但是我们还得更新列号
        col += len(r.text)

      pending_contents.commit()
      pending_content = pending_contents.contents
      locations = pending_contents.locations
      del pending_contents

      # 然后我们再根据现在的情况把内容写回去
      # 先检查前面的特殊块项
      if len(pending_content) > 0:
        # 如果这段是特殊块，我们把所有相邻的文本内容都合并了，这样省得后面再合并
        if special_block_reason is not None:
          newcontents = []
          newlocations = []
          s = ''
          loc = -1
          for i in range(0, len(pending_content)): # pylint: disable=consider-using-enumerate
            content = pending_content[i]
            if isinstance(content, (StringLiteral, TextFragmentLiteral)):
              if loc == -1:
                # 开启新串
                s = content.get_string()
                loc = locations[i]
              else:
                s += content.get_string()
            else:
              # 结束一个串
              if loc != -1:
                newcontents.append(StringLiteral.get(s, self.context))
                newlocations.append(loc)
              newcontents.append(content)
              newlocations.append(locations[i])
              s = ''
              loc = -1
          if loc != -1:
            newcontents.append(StringLiteral.get(s, self.context))
            newlocations.append(loc)
          pending_content = newcontents
          locations = newlocations

        def get_dest_block():
          nonlocal current_dest_block
          if current_dest_block is not None:
            return current_dest_block
          if len(current_ongoing_list) > 0:
            current_dest_block = current_ongoing_list[-1].add_list_item().create_block()
          else:
            current_dest_block = dest.create_block()
          return current_dest_block

        # 开始写回主循环
        # 如果正常模式下我们有连续的字符串、文本片段，我们使用同一个 IMElementOp
        last_element : IMElementOp | None = None
        for i in range(0, len(pending_content)): # pylint: disable=consider-using-enumerate
          content = pending_content[i]
          col = locations[i]
          if isinstance(content, (StringLiteral, TextFragmentLiteral)):
            if special_block_reason is not None:
              # 特殊块内容
              assert isinstance(content, StringLiteral)
              if i == 0 and last_special_block is not None:
                # 可以并入前面的特殊块
                last_special_block.content.add_operand(content)
              else:
                # 需要新建
                last_special_block = IMSpecialBlockOp.create(content, special_block_reason, '', self.context.get_DILocation(self.difile, page, row, col))
                get_dest_block().push_back(last_special_block)
            else:
              if last_element is None:
                last_element = IMElementOp.create(content, '', self.context.get_DILocation(self.difile, page, row, col))
                get_dest_block().push_back(last_element)
              else:
                last_element.content.add_operand(content)
          elif isinstance(content, AssetData):
            last_special_block = None
            last_element = None
            element = IMElementOp.create(content, '', self.context.get_DILocation(self.difile, page, row, col))
            get_dest_block().push_back(element)
          elif isinstance(content, Operation):
            # 应该是 ErrorOp
            get_dest_block().push_back(content)
          else:
            raise PPInternalError("Unexpected content type")
        # 有内容的情况处理完毕
      else:
        # 无内容的情况，空段落
        last_special_block = None

      # 写回完成，该段落处理完毕
    def handle_table(t : docx.table.Table):
      nonlocal page
      nonlocal row
      nonlocal last_special_block
      nonlocal current_ongoing_list

      col = 1
      last_special_block = None
      current_ongoing_list.clear()
      nrows = len(t.rows)
      ncols = len(t.columns)
      result = IMTableOp.create(rowcount=nrows, columncount=ncols, name='', loc=self.context.get_DILocation(self.difile, page, row, col))
      dest.create_block().push_back(result)
      errors : list[ErrorOp] = []
      for x in range(nrows):
        row += 1
        col = 0
        for y in range(ncols):
          col += 1
          cellsrc = t.cell(x, y)
          if len(cellsrc.paragraphs) > 0:
            contents = _DocxTextCoalescer(self.context)
            is_first_paragraph = True
            for p in cellsrc.paragraphs:
              if not is_first_paragraph:
                # 换段落的话加一个 '\n' 到字符串值中
                contents.add_value(StringLiteral.get("\n", self.context), col)
              else:
                is_first_paragraph = False
              for r in p.runs:
                res = self.parse_run(r, page, row, col)
                if isinstance(res, Value):
                  contents.add_value(res, col)
                elif isinstance(res, ErrorOp):
                  errors.append(res)
            contents.commit()
            celldest = result.get_cell_operand(x, y)
            for v in contents.contents:
              celldest.add_operand(v)
      if len(errors) > 0:
        for e in reversed(errors):
          e.insert_before(result)
      # 该表格处理完毕
    for pt in self.iter_block_items(self.dochandle):
      if isinstance(pt, docx.text.paragraph.Paragraph):
        handle_paragraph(pt)
      elif isinstance(pt, docx.table.Table):
        handle_table(pt)
      else:
        raise PPInternalError("Unexpected block item type")
    # 所有段落处理完毕，结束

  def parse_run(self, r : docx.text.run.Run, page : int, row : int, col : int) -> ErrorOp | AssetData | TextFragmentLiteral | StringLiteral | None:
    if len(r.text) > 0:
      # 文本内容
      s = StringLiteral.get(r.text, self.context)
      style = self.get_character_style(r)
      if isinstance(style, TextStyleLiteral):
        return TextFragmentLiteral.get(self.context, s, style)
      elif style is None:
        return s
      # 这是被划掉的内容，不产出 Value
      return None
    else:
      # 内嵌内容
      # 参考 parse_rels() 中的链接
      if "Graphic" in r._r.xml:
        for rid in self.rels:
          if rid in r._r.xml:
            href = self.get_path_from_rid(rid)
            if href in self.ziphandle.namelist():
              data = self.ziphandle.read(href)
              mime, encoding = mimetypes.guess_type(href)
              loc = self.context.get_DILocation(self.difile, page, row, col)
              if mime is None:
                msg = "Unknown media type for media \"" + href + "\""
                MessageHandler.critical_warning(msg, self.filepath)
                textstr = StringLiteral.get(href, self.context)
                msgstr = StringLiteral.get(msg, self.context)
                return IMErrorElementOp.create(name = '', loc = loc, content = textstr, error_code='docx-bad-media-ref', error_msg = msgstr)
              fullpath = self.get_full_path_from_epath(href)
              if fmt := ImageAssetData.get_format_from_mime_type(mime):
                value = self.context.create_image_asset_data_embedded(fullpath, data, fmt)
              elif fmt := AudioAssetData.get_format_from_mime_type(mime):
                value = self.context.create_audio_asset_data_embedded(fullpath, data, fmt)
              else:
                textstr = StringLiteral.get(href, self.context)
                msgstr = StringLiteral.get(mime, self.context)
                return IMErrorElementOp.create(name = '', loc = loc, content = textstr, error_code='docx-bad-media-ref', error_msg = msgstr)
              if value is None:
                msg = "Cannot resolve reference to media \"" + href + "\" (please check file presence or security violation)"
                MessageHandler.critical_warning(msg, self.filepath)
                textstr = StringLiteral.get(href, self.context)
                msgstr = StringLiteral.get(msg, self.context)
                return IMErrorElementOp.create(name = '', loc = loc, content = textstr, error_code='docx-bad-media-ref', error_msg = msgstr)

              assert isinstance(value, AssetData)
              return value
    # 到这的话说明引用的内容暂不支持，直接跳过
    return None

class _DocxTextCoalescer:
  contents : list[Value | Operation]
  locations : list[int]
  s : str
  style : TextStyleLiteral | None
  loc : int
  context : Context

  def __init__(self, ctx : Context) -> None:
    self.contents = []
    self.locations = []
    self.s = ''
    self.style = None
    self.loc = -1
    self.context = ctx

  def commit(self):
    if len(self.s) > 0:
      v = StringLiteral.get(self.s, self.context)
      if self.style is not None:
        v = TextFragmentLiteral.get(self.context, v, self.style)
      self.contents.append(v)
      self.locations.append(self.loc)
    self.s = ''
    self.style = None
    self.loc = -1

  def add_value(self, content : Value | Operation, location : int):
    if isinstance(content, (StringLiteral, TextFragmentLiteral)):
      cur_s = content.get_string()
      cur_style = None
      if isinstance(content, TextFragmentLiteral):
        cur_style = content.style
      if self.loc == -1:
        # 开启新串
        self.s = cur_s
        self.style = cur_style
        self.loc = location
      elif self.style is cur_style:
        self.s += cur_s
      else:
        # 不同的串
        self.commit()
        self.s = cur_s
        self.style = cur_style
        self.loc = location
    else:
      self.commit()
      self.contents.append(content)
      self.locations.append(location)


@FrontendDecl('docx', input_decl=IODecl('OfficeOpenXML files', match_suffix=('docx',), nargs='+'), output_decl=IMDocumentOp)
class ReadDOCX(TransformBase):
  _ctx : Context

  def __init__(self, _ctx: Context) -> None:
    super().__init__(_ctx)
    self._ctx = _ctx

  def run(self) -> IMDocumentOp | typing.List[IMDocumentOp]:
    if len(self.inputs) == 1:
      return _DOCXParseContext.parse_docx(self._ctx, self.inputs[0])
    results = []
    for f in self.inputs:
      results.append(_DOCXParseContext.parse_docx(self._ctx, f))
    return results
